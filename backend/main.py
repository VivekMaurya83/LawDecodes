import json
import os
import shutil
import uuid
import sys
from pathlib import Path
from typing import List, Dict, Any
from fastapi import FastAPI, HTTPException, Depends, UploadFile, File, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from fastapi.security import OAuth2PasswordRequestForm
from pydantic import BaseModel
from starlette.responses import FileResponse
from dotenv import load_dotenv
from fastapi.responses import StreamingResponse
from docx import Document
from docx.shared import Inches
import io

# --- Load Environment Variables ---
# This line MUST be at the very top of your application's entry point.
load_dotenv()

# --- Add Project Subdirectories to Python Path ---
BASE_DIR = Path(__file__).resolve().parent.parent
EXTRACTION_PIPELINE_DIR = BASE_DIR / "Extraction_Pipeline"
RAG_MODEL_DIR = BASE_DIR / "RAG Model"
sys.path.append(str(EXTRACTION_PIPELINE_DIR))
sys.path.append(str(RAG_MODEL_DIR))

# --- Import Pipeline and RAG Components ---
from main_flow import initial_text_extraction, run_summarization_task
# NOTE: Ensure your RAG project's entry point is named rag_chatbot.py
from rag_chatbot import LegalRAGChatbot
from jargons import process_and_extract_jargon

# --- Import from your other local files ---
from db import (
    create_user, find_user_by_username, get_profile_details, 
    update_profile_details, get_user_chats, save_user_chat, 
    delete_user_chat, SECURITY_QUESTIONS
)
from auth import create_access_token, get_current_user, verify_password

# --- Directory Setup ---
UPLOAD_DIRECTORY = BASE_DIR / "backend" / "uploads"
SUMMARIZED_FORM_DIR = EXTRACTION_PIPELINE_DIR / "summarized_form"
os.makedirs(UPLOAD_DIRECTORY, exist_ok=True)
os.makedirs(SUMMARIZED_FORM_DIR, exist_ok=True)

app = FastAPI()

# --- CORS Middleware ---
app.add_middleware(
    CORSMiddleware, 
    allow_origins=["*"], 
    allow_credentials=True, 
    allow_methods=["*"], 
    allow_headers=["*"]
)

# --- RAG Cache: Stores prepared chatbot instances ---
RAG_CHATBOTS: Dict[str, LegalRAGChatbot] = {}

# --- Pydantic Models ---
class SignupRequest(BaseModel):
    username: str; email: str; password: str; securityQuestion: int; securityAnswer: str
class ProfileUpdateRequest(BaseModel):
    name: str; email: str
class Chat(BaseModel):
    username: str; chat_id: str; title: str; html_content: str = ""
class ChatSaveRequest(BaseModel):
    chat_id: str; title: str; html_content: str
class SummarizeRequest(BaseModel):
    filePath: str
class RAGQueryRequest(BaseModel):
    extractedFilePath: str
    query: str
class TextUploadRequest(BaseModel):
    raw_text: str


class VerifySecurityRequest(BaseModel):
    username: str
    securityAnswer: str


def _prepare_rag_task(full_extracted_path: str):
    """Background task for RAG preparation ONLY."""
    print(f"BACKGROUND TASK (RAG): Starting preparation for {full_extracted_path}")
    try:
        chatbot_instance = LegalRAGChatbot(documents_path=full_extracted_path)
        RAG_CHATBOTS[full_extracted_path] = chatbot_instance
        print(f"BACKGROUND TASK (RAG): SUCCESS! RAG is ready.")
    except Exception as e:
        print(f"BACKGROUND TASK (RAG): FAILED. Error: {e}", file=sys.stderr)
        RAG_CHATBOTS[full_extracted_path] = None

def _extract_jargon_task(full_extracted_path: str, jargon_output_path: str):
    """Background task for Jargon extraction ONLY."""
    print(f"BACKGROUND TASK (Jargon): Starting extraction for {full_extracted_path}")
    try:
        dict_path = str(BASE_DIR / "Extraction_Pipeline" / "legal_dict.json")
        jargon_data = process_and_extract_jargon(full_extracted_path, dict_path)
        with open(jargon_output_path, 'w', encoding='utf-8') as f:
            json.dump(jargon_data, f, indent=2)
        print(f"BACKGROUND TASK (Jargon): SUCCESS! Jargon file created.")
    except Exception as e:
        print(f"BACKGROUND TASK (Jargon): FAILED. Error: {e}", file=sys.stderr)


# --- AUTH, SIGNUP, & PROFILE ENDPOINTS ---
@app.post("/login")
def login(form_data: OAuth2PasswordRequestForm = Depends()):
    user = find_user_by_username(form_data.username)
    if not user or not verify_password(form_data.password, user.get('password')):
        raise HTTPException(status_code=401, detail="Incorrect username or password", headers={"WWW-Authenticate": "Bearer"})
    access_token = create_access_token(data={"sub": user.get('username')})
    return {"access_token": access_token, "token_type": "bearer"}

@app.post("/signup")
def signup(data: SignupRequest):
    result = create_user(data.username.strip(), data.email.strip(), data.password.strip(), data.securityQuestion, data.securityAnswer.strip().lower())
    if not result["success"]: raise HTTPException(status_code=400, detail=result["message"])
    return result

@app.get("/api/profile")
def profile(current_user: str = Depends(get_current_user)):
    profile_details = get_profile_details(current_user)
    if not profile_details: raise HTTPException(status_code=404, detail="User not found")
    return profile_details

@app.post("/api/upload-text")
def upload_raw_text(data: TextUploadRequest, current_user: str = Depends(get_current_user)):
    """
    Receives raw text, saves it to a unique .txt file, and returns the path.
    This allows the 'paste text' feature to use the same workflow as file uploads.
    """
    try:
        unique_filename = f"{uuid.uuid4()}.txt"
        file_path = UPLOAD_DIRECTORY / unique_filename
        
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(data.raw_text)
            
        return {
            "filePath": f"backend/uploads/{unique_filename}",
            "originalFilename": "pastedText.txt"
        }
    except Exception as e:
        print(f"ERROR in /api/upload-text endpoint: {e}", file=sys.stderr)
        raise HTTPException(status_code=500, detail=f"Could not process text: {e}")

@app.get("/security-questions")
def get_security_questions():
    return {"questions": SECURITY_QUESTIONS}


@app.get("/get-security-question/{username}")
def get_security_question_for_user(username: str):
    """Return the security question text for a given username (case-insensitive)."""
    print(f"[GET /get-security-question] username={username}")
    user = find_user_by_username(username)
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    qid = user.get("security_question_id")
    question = SECURITY_QUESTIONS.get(qid, None)
    if not question:
        raise HTTPException(status_code=404, detail="Security question not set for this user")
    print(f"[GET /get-security-question] found question id={qid}")
    return {"securityQuestion": question}


@app.post("/verify-security")
def verify_security_answer(data: VerifySecurityRequest):
    """Verify a user's answer to their security question.

    Returns 200 with {"verified": True} on success, 401 on incorrect answer, 404 if user missing.
    """
    print(f"[POST /verify-security] username={data.username} payload_answer={(data.securityAnswer or '')[:60]}")
    user = find_user_by_username(data.username)
    if not user:
        raise HTTPException(status_code=404, detail="User not found")

    stored_answer = user.get("security_answer", "").strip().lower()
    provided = (data.securityAnswer or "").strip().lower()
    if stored_answer and provided and stored_answer == provided:
        print(f"[POST /verify-security] verified OK for user={data.username}")
        # Create an access token so the frontend can log the user in after verification
        try:
            access_token = create_access_token(data={"sub": user.get('username')})
            return {"verified": True, "access_token": access_token, "token_type": "bearer"}
        except Exception as e:
            print(f"[POST /verify-security] error creating token: {e}", file=sys.stderr)
            return {"verified": True}
    raise HTTPException(status_code=401, detail="Incorrect security answer")

# --- CHAT HISTORY ENDPOINTS ---
@app.get("/api/chats", response_model=List[Chat])
def chats(current_user: str = Depends(get_current_user)):
    return get_user_chats(current_user)

@app.post("/api/chats")
def save_chat(data: ChatSaveRequest, current_user: str = Depends(get_current_user)):
    if not save_user_chat(current_user, data.chat_id, data.title, data.html_content)["success"]:
        raise HTTPException(status_code=500, detail="Could not save chat.")
    return {"message": "Chat saved successfully"}

@app.delete("/api/chats/{chat_id}")
def delete_chat(chat_id: str, current_user: str = Depends(get_current_user)):
    if not delete_user_chat(current_user, chat_id):
        raise HTTPException(status_code=404, detail="Chat not found or could not be deleted.")
    return {"message": "Chat deleted successfully."}



# --- FILE & UNIFIED PROCESSING ENDPOINTS ---
@app.post("/api/upload")
def upload_file(file: UploadFile = File(...), current_user: str = Depends(get_current_user)):
    try:
        file_extension = os.path.splitext(file.filename)[1]
        unique_filename = f"{uuid.uuid4()}{file_extension}"
        file_path = UPLOAD_DIRECTORY / unique_filename
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        return {"filePath": f"backend/uploads/{unique_filename}", "originalFilename": file.filename}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Could not upload file: {e}")

# --- MODIFIED & NEW ENDPOINTS ---

# In main.py

@app.post("/api/summarize")
def summarize_document(data: SummarizeRequest, background_tasks: BackgroundTasks):
    """
    Performs fast text extraction immediately, then schedules Summarization,
    RAG Prep, and Jargon Extraction to run in PARALLEL in the background.
    """
    document_path = str(BASE_DIR / data.filePath)
    if not os.path.exists(document_path):
        raise HTTPException(status_code=404, detail="Document to be summarized not found.")

    # --- Define all output paths ---
    base_name = os.path.splitext(os.path.basename(document_path))[0]
    report_filename = f"{base_name}_ANALYSIS_REPORT.txt"
    extracted_filename = f"{base_name}_extracted.txt"
    jargon_filename = f"{base_name}_jargons.json"

    report_full_path = str(SUMMARIZED_FORM_DIR / report_filename)
    extracted_full_path = str(SUMMARIZED_FORM_DIR / extracted_filename)
    jargon_full_path = str(SUMMARIZED_FORM_DIR / jargon_filename)

    relative_extracted_path = f"Extraction_Pipeline/summarized_form/{extracted_filename}"

    # --- Step 1: Perform FAST, UNIFIED text extraction SYNCHRONOUSLY ---
    print(f"IMMEDIATE TASK: Starting initial text extraction for {document_path}")

    # --- THE FIX IS HERE ---
    # REMOVE the special 'if file_extension == ".txt"' check.
    # By calling initial_text_extraction for all file types, you ensure that the
    # text from the paste feature undergoes the same crucial cleaning/standardization
    # (read and rewrite) that your main_flow.py already specifies.
    # This consistency is what the RAG model requires.
    try:
        print(f"Using unified extraction pipeline for '{os.path.basename(document_path)}' to ensure consistency.")
        
        # This single call now correctly handles PDF, DOCX, and TXT files from the paste feature.
        success = initial_text_extraction(document_path, extracted_full_path)

    except Exception as e:
        print(f"ERROR during the unified text extraction step: {e}", file=sys.stderr)
        raise HTTPException(status_code=500, detail=f"Failed to process the document file: {e}")

    if not success:
        raise HTTPException(status_code=500, detail="Failed to perform initial text extraction from the document.")

    # Verification step to ensure the extracted file is not empty.
    if not os.path.exists(extracted_full_path) or os.path.getsize(extracted_full_path) == 0:
        raise HTTPException(status_code=500, detail="Text extraction resulted in an empty file. Cannot proceed.")

    print("IMMEDIATE TASK: Initial text extraction successful. The file is now clean for all background tasks.")

    # --- Step 2: Schedule the THREE long-running tasks to run IN PARALLEL ---
    background_tasks.add_task(run_summarization_task, extracted_full_path, report_full_path)
    background_tasks.add_task(_prepare_rag_task, extracted_full_path)
    background_tasks.add_task(_extract_jargon_task, extracted_full_path, jargon_full_path)

    print("All three background tasks (Summarize, RAG, Jargon) have been scheduled.")

    return {
        "message": "Full document processing started in parallel.",
        "report_filename": report_filename,
        "extractedFilePath": relative_extracted_path,
        "jargonFilename": jargon_filename
    }

# NEW ENDPOINT to fetch jargon data
@app.get("/api/get-jargons/{jargon_filename}")
def get_jargon_data(jargon_filename: str, current_user: str = Depends(get_current_user)):
    jargon_path = SUMMARIZED_FORM_DIR / jargon_filename
    if not jargon_path.is_file():
        raise HTTPException(status_code=404, detail="Jargon file not yet available or does not exist.")
    
    with open(jargon_path, 'r', encoding='utf-8') as f:
        jargon_data = json.load(f)
        
    return JSONResponse(content=jargon_data)

@app.get("/api/get-report/{report_filename}")
def get_summary_report(report_filename: str, current_user: str = Depends(get_current_user)):
    report_path = SUMMARIZED_FORM_DIR / report_filename
    if not report_path.is_file():
        raise HTTPException(status_code=404, detail="Report not yet available or does not exist.")
    with open(report_path, "r", encoding="utf-8") as f:
        report_content = f.read()
    return {"summary": report_content}

# --- THIS IS THE NEW, ROBUST CODE ---
@app.post("/api/ask-rag", response_model=Dict[str, Any])
def ask_rag_question(data: RAGQueryRequest, current_user: str = Depends(get_current_user)):
    """
    Asks a question to a RAG chatbot.
    If the chatbot isn't in memory, it initializes it on-demand.
    """
    if not data.query:
        raise HTTPException(status_code=400, detail="Query cannot be empty.")

    full_extracted_path = str(BASE_DIR / data.extractedFilePath)

    # Check if the physical file for Q&A even exists.
    if not os.path.exists(full_extracted_path):
        raise HTTPException(status_code=404, detail="The source document for Q&A cannot be found on the server.")

    # --- THE LAZY INITIALIZATION LOGIC ---
    # Check if the chatbot is in our in-memory cache.
    if full_extracted_path not in RAG_CHATBOTS:
        print(f"RAG model for '{data.extractedFilePath}' not in memory. Initializing on-demand...")
        try:
            # If not, create a new instance and add it to the cache.
            chatbot_instance = LegalRAGChatbot(documents_path=full_extracted_path)
            RAG_CHATBOTS[full_extracted_path] = chatbot_instance
            print("On-demand initialization successful.")
        except Exception as e:
            print(f"On-demand RAG initialization FAILED. Error: {e}", file=sys.stderr)
            raise HTTPException(status_code=500, detail=f"Could not prepare the document for Q&A: {e}")

    # Now we are guaranteed to have a chatbot instance.
    chatbot = RAG_CHATBOTS.get(full_extracted_path)

    # This check handles the rare case where initialization failed silently.
    if chatbot is None:
        raise HTTPException(status_code=500, detail="Q&A session for this document is in a failed state. Please try reprocessing.")

    print(f"Querying RAG for '{data.extractedFilePath}' with question: '{data.query}'")
    response = chatbot.chat(data.query)
    return response


# --- NEW ENDPOINT FOR SIDE-BY-SIDE COMPARISON ---
@app.get("/api/get-extracted-text/{extracted_filename}")
def get_extracted_text(extracted_filename: str, current_user: str = Depends(get_current_user)):
    """
    (SECURED) Fetches the raw extracted text content from a specified file.
    This is used to populate the 'Original' side of the comparison view.
    """
    text_path = SUMMARIZED_FORM_DIR / extracted_filename
    if not text_path.is_file():
        raise HTTPException(status_code=404, detail="Extracted text file not found. It may still be processing.")
    
    with open(text_path, "r", encoding="utf-8") as f:
        text_content = f.read()
        
    return {"original_text": text_content}

# --- NEW ENDPOINT FOR EXPORTING DOCX ---

@app.get("/api/export-docx/{report_filename}/{extracted_filename}")
async def export_summary_as_docx(report_filename: str, extracted_filename: str, current_user: str = Depends(get_current_user)):
    """
    (SECURED) Generates a .docx file containing the summary and original text
    and streams it back to the user for download.
    """
    report_path = SUMMARIZED_FORM_DIR / report_filename
    extracted_text_path = SUMMARIZED_FORM_DIR / extracted_filename

    if not report_path.is_file() or not extracted_text_path.is_file():
        raise HTTPException(status_code=404, detail="Required summary or text files not found.")

    try:
        with open(report_path, "r", encoding="utf-8") as f:
            summary_content = f.read()
        with open(extracted_text_path, "r", encoding="utf-8") as f:
            original_content = f.read()
            
        # Create a new Word document
        doc = Document()
        doc.add_heading('Legal Document Analysis Report', level=0)

        # --- Add Summary Section ---
        doc.add_heading('Generated Summary', level=1)
        # Split summary by lines to handle headers and paragraphs
        for line in summary_content.split('\n'):
            if line.startswith('# '): # Main headers in markdown
                # Add as a heading, removing the '#'
                doc.add_heading(line.lstrip('# ').strip(), level=2)
            elif line.strip() == '---':
                doc.add_page_break()
            elif line.strip(): # Add non-empty lines as paragraphs
                doc.add_paragraph(line)

        # Save the document to an in-memory stream
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0) # Move cursor to the beginning of the stream

        # Define headers for the file download
        headers = {
            'Content-Disposition': f'attachment; filename="Analysis_Report_{report_filename.replace("_ANALYSIS_REPORT.txt", "")}.docx"'
        }
        
        return StreamingResponse(file_stream, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers=headers)

    except Exception as e:
        print(f"Error creating DOCX: {e}", file=sys.stderr)
        raise HTTPException(status_code=500, detail="Failed to generate the DOCX report.")